import glob
from docx import Document
from docx.shared import Inches

document = Document('demo.docx')
p = document.add_paragraph()
r = p.add_run()
i = 1


data = "C:\qrcode\qr_\qr_388.png"
file_name = data.split('\\')[3]
input_name = file_name.split('.')[0]

r.add_picture('C:\qrcode\qr_\\'+file_name, width=Inches(0.75))
r.add_text('   '+input_name+'       ')

#이부분은 대량으로 뽑을때
#file_list = glob.glob('C:\qrcode\pr_\*.*')
#for data in file_list:
#    enter = i % 5
#    file_name = data.split('\\')[3]
#    input_name = file_name.split('.')[0]
#    if enter == 0:
#        r.add_picture('C:\qrcode\pp_\\'+file_name, width=Inches(0.75))
#        r.add_text('   '+input_name+ " \n" )
#    else :
#        r.add_picture('C:\qrcode\pp_\\'+file_name, width=Inches(0.75))
#        r.add_text('   '+input_name+'       ')
        #print(data.split('\\')[3])
#    i = i + 1

document.save('demo.docx')
